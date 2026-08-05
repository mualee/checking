"""Generate the Lao-language Word report (python-docx), returned as in-memory bytes.

Font handling: python-docx's run.font.name only sets w:ascii. For a complex script
like Lao, Word may substitute a font for the eastAsia/complex-script runs unless we
also set w:hAnsi, w:eastAsia, and w:cs. We therefore set all four via direct rFonts
XML through _set_run_font(). (The font need not be installed on the server — the
.docx is valid regardless; installation only affects viewer-side rendering.)
"""
from __future__ import annotations

import io
from decimal import Decimal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from app.core.constants import TOTAL_LABEL
from app.services.calculations import Table1Row, Table2Row, Table3Group, ValidationResult
from app.utils.money import format_lao_amount

REPORT_FONT = "Phetsarath OT"
HEADER_FILL = "D9D9D9"   # light gray
TOTAL_FILL = "FFF2CC"    # light amber, distinct from header

TABLE1_HEADERS = ["ເດືອນ", "ໜີ້", "ມີ", "ສ່ວນຕ່າງ", "ຍອດເຫຼືອທ້າຍເດືອນ"]
TABLE2_HEADERS = ["ເດືອນ", "ໜີ້", "ມີ", "ສ່ວນຕ່າງ"]
TABLE3_HEADERS = ["ເດືອນ", "ເລກທີ", "ໜີ້", "ລາຍລະອຽດ"]


def _set_run_font(run, name: str = REPORT_FONT, *, size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), name)


def _set_cell_shading(cell, hex_fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _write_cell(cell, text: str, *, bold: bool = False, align=None) -> None:
    para = cell.paragraphs[0]
    for existing in list(para.runs):  # ensure exactly one styled run
        existing._element.getparent().remove(existing._element)
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    _set_run_font(run, bold=bold)


def _add_paragraph(doc: Document, text: str, *, size: int = 11, bold: bool = False, align=None):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    _set_run_font(run, size=size, bold=bold)
    return para


def _amt(d: Decimal) -> str:
    return format_lao_amount(d)


def _build_table1(doc: Document, table1: list[Table1Row]) -> None:
    table = doc.add_table(rows=1, cols=len(TABLE1_HEADERS))
    table.style = "Table Grid"
    for i, h in enumerate(TABLE1_HEADERS):
        _write_cell(table.rows[0].cells[i], h, bold=True)
        _set_cell_shading(table.rows[0].cells[i], HEADER_FILL)
    for r in table1:
        is_total = r.month == TOTAL_LABEL
        cells = table.add_row().cells
        _write_cell(cells[0], r.month, bold=is_total)
        _write_cell(cells[1], _amt(r.debit), bold=is_total, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _write_cell(cells[2], _amt(r.credit), bold=is_total, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _write_cell(cells[3], _amt(r.diff), bold=is_total, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _write_cell(cells[4], _amt(r.end_balance), bold=is_total, align=WD_ALIGN_PARAGRAPH.RIGHT)
        if is_total:
            for c in cells:
                _set_cell_shading(c, TOTAL_FILL)


def _build_table2(doc: Document, table2: list[Table2Row]) -> None:
    table = doc.add_table(rows=1, cols=len(TABLE2_HEADERS))
    table.style = "Table Grid"
    for i, h in enumerate(TABLE2_HEADERS):
        _write_cell(table.rows[0].cells[i], h, bold=True)
        _set_cell_shading(table.rows[0].cells[i], HEADER_FILL)
    for r in table2:
        cells = table.add_row().cells
        _write_cell(cells[0], r.month)
        _write_cell(cells[1], _amt(r.debit), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _write_cell(cells[2], _amt(r.credit), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _write_cell(cells[3], _amt(r.diff), align=WD_ALIGN_PARAGRAPH.RIGHT)


def _build_table3(doc: Document, groups: list[Table3Group]) -> None:
    if not groups:
        _add_paragraph(doc, "— ບໍ່ພົບລາຍຈ່າຍປະຈຳທີ່ຕິດຕໍ່ກັນ —", size=10)
        return
    for g in groups:
        _add_paragraph(
            doc,
            f"ຈຳນວນ {_amt(g.amount)} — ຕິດຕໍ່ກັນ {g.month_count} ເດືອນ",
            size=11, bold=True,
        )
        table = doc.add_table(rows=1, cols=len(TABLE3_HEADERS))
        table.style = "Table Grid"
        for i, h in enumerate(TABLE3_HEADERS):
            _write_cell(table.rows[0].cells[i], h, bold=True)
            _set_cell_shading(table.rows[0].cells[i], HEADER_FILL)
        for r in g.rows:
            cells = table.add_row().cells
            _write_cell(cells[0], r.month)
            _write_cell(cells[1], r.txn_number)
            _write_cell(cells[2], _amt(r.debit), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _write_cell(cells[3], r.description)
        doc.add_paragraph()


def generate_statement_report(
    *,
    customer_name: str,
    account_no: str,
    period: str,
    table1: list[Table1Row],
    table2: list[Table2Row],
    table3: list[Table3Group] | None = None,
    validation: ValidationResult | None = None,
) -> bytes:
    """Build the .docx entirely in memory and return raw bytes (never touches disk)."""
    doc = Document()

    _add_paragraph(doc, "ບົດລາຍງານການກວດສອບ Statement", size=16, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, f"ລູກຄ້າ: {customer_name}", size=11)
    _add_paragraph(doc, f"ເລກບັນຊີ: {account_no}", size=11)
    _add_paragraph(doc, f"ໄລຍະ: {period}", size=11)

    _add_paragraph(doc, "ຕາຕະລາງ 1: ສະຫຼຸບຕໍ່ເດືອນ", size=13, bold=True)
    _build_table1(doc, table1)

    doc.add_paragraph()
    _add_paragraph(doc, "ຕາຕະລາງ 2: 3 ເດືອນທຸລະກຳຮັບສູງສຸດ", size=13, bold=True)
    _build_table2(doc, table2)

    doc.add_paragraph()
    _add_paragraph(doc, "ຕາຕະລາງ 3: ລາຍຈ່າຍປະຈຳ (ຕິດຕໍ່ກັນ 3 ເດືອນຂຶ້ນໄປ)", size=13, bold=True)
    _build_table3(doc, table3 or [])

    if validation is not None:
        doc.add_paragraph()
        status = "ຖືກຕ້ອງ (matched)" if validation.matched else \
                 f"ພົບຄວາມຄາດເຄື່ອນ {validation.mismatch_count} ແຖວ"
        _add_paragraph(doc, f"ຜົນກວດສອບຄວາມຖືກຕ້ອງ: {status}", size=11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
