"""Structural assertions on the generated .docx — no emulator required."""
from __future__ import annotations

import io
from decimal import Decimal

from docx import Document
from docx.oxml.ns import qn

from app.core.constants import TOTAL_LABEL
from app.services.calculations import (
    build_table1,
    build_table2,
    build_table3,
    validate_balances,
)
from app.services.docx_report import (
    HEADER_FILL,
    REPORT_FONT,
    TOTAL_FILL,
    generate_statement_report,
)
from app.models.transaction import TransactionRow


def _sample_tables():
    txns = [
        TransactionRow(date="01/05/2024", txn_number="B1", description="d",
                       debit=Decimal("0"), credit=Decimal("5000000"), balance=Decimal("6000000")),
        TransactionRow(date="10/05/2024", txn_number="B2", description="d",
                       debit=Decimal("2000000"), credit=Decimal("0"), balance=Decimal("4000000")),
        TransactionRow(date="05/06/2024", txn_number="B3", description="d",
                       debit=Decimal("0"), credit=Decimal("8000000"), balance=Decimal("12000000")),
        TransactionRow(date="05/07/2024", txn_number="B4", description="d",
                       debit=Decimal("0"), credit=Decimal("9000000"), balance=Decimal("21000000")),
    ]
    return build_table1(txns), build_table2(txns)


def _generate() -> Document:
    t1, t2 = _sample_tables()
    data = generate_statement_report(
        customer_name="ທ້າວ ສົມໃຈ", account_no="001-234-567",
        period="05/2024 - 07/2024", table1=t1, table2=t2,
    )
    assert data[:2] == b"PK", "output must be a valid .docx (zip) file"
    return Document(io.BytesIO(data))


def test_document_has_two_tables():
    doc = _generate()
    assert len(doc.tables) == 2


def test_table_dimensions():
    doc = _generate()
    t1, t2 = _sample_tables()
    # header row + data rows
    assert len(doc.tables[0].rows) == len(t1) + 1
    assert len(doc.tables[0].columns) == 5
    assert len(doc.tables[1].rows) == len(t2) + 1
    assert len(doc.tables[1].columns) == 4


def _cell_fill(cell) -> str | None:
    tcpr = cell._tc.find(qn("w:tcPr"))
    if tcpr is None:
        return None
    shd = tcpr.find(qn("w:shd"))
    return shd.get(qn("w:fill")) if shd is not None else None


def test_header_row_shaded():
    doc = _generate()
    for cell in doc.tables[0].rows[0].cells:
        assert _cell_fill(cell) == HEADER_FILL


def test_total_row_highlighted_and_distinct_from_header():
    doc = _generate()
    total_row = doc.tables[0].rows[-1]
    assert total_row.cells[0].text == TOTAL_LABEL
    for cell in total_row.cells:
        assert _cell_fill(cell) == TOTAL_FILL
    assert TOTAL_FILL != HEADER_FILL


def test_table2_has_no_total_row():
    doc = _generate()
    # No cell in table 2 should carry the total label.
    for row in doc.tables[1].rows:
        assert row.cells[0].text != TOTAL_LABEL


def test_runs_use_phetsarath_font_with_full_rfonts():
    doc = _generate()
    # Sample a header run and check all four rFonts attributes are set.
    run = doc.tables[0].rows[0].cells[0].paragraphs[0].runs[0]
    assert run.font.name == REPORT_FONT
    rfonts = run._element.get_or_add_rPr().find(qn("w:rFonts"))
    assert rfonts is not None
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        assert rfonts.get(qn(attr)) == REPORT_FONT


def test_table3_recurring_section_renders():
    txns = [
        TransactionRow(date="05/05/2024", txn_number="R1", description="LOAN",
                       debit=Decimal("500000"), credit=Decimal("0"), balance=Decimal("1")),
        TransactionRow(date="05/06/2024", txn_number="R2", description="LOAN",
                       debit=Decimal("500000"), credit=Decimal("0"), balance=Decimal("1")),
        TransactionRow(date="05/07/2024", txn_number="R3", description="LOAN",
                       debit=Decimal("500000"), credit=Decimal("0"), balance=Decimal("1")),
    ]
    t1, t2, t3 = build_table1(txns), build_table2(txns), build_table3(txns)
    assert len(t3) == 1  # one recurring group detected
    data = generate_statement_report(
        customer_name="X", account_no="1", period="p",
        table1=t1, table2=t2, table3=t3,
    )
    doc = Document(io.BytesIO(data))
    # Tables 1 & 2 plus one Table-3 group table.
    assert len(doc.tables) == 3
    t3_table = doc.tables[2]
    assert len(t3_table.columns) == 4  # month, no, debit, description
    assert len(t3_table.rows) == 1 + 3  # header + 3 months


def test_validation_line_rendered_when_provided():
    txns = [
        TransactionRow(date="01/05/2024", txn_number="B1", description="d",
                       debit=Decimal("0"), credit=Decimal("100"), balance=Decimal("1100")),
    ]
    t1, t2 = build_table1(txns), build_table2(txns)
    val = validate_balances(txns, Decimal("1000"))
    data = generate_statement_report(
        customer_name="X", account_no="1", period="p",
        table1=t1, table2=t2, validation=val,
    )
    doc = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "matched" in text
